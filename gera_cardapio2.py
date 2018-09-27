import docx
import re
from classes import Refeicao, Semana
import datetime
import os
from jinja2 import Environment, FileSystemLoader

PATH = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_ENVIRONMENT = Environment(
    autoescape=False,
    loader=FileSystemLoader(os.path.join(PATH, 'templates')),
    trim_blocks=False)

arrayCardapio = []
final = []

def render_template(template_filename, context):
    return TEMPLATE_ENVIRONMENT.get_template(template_filename).render(context)

def adicionaDias(cardapio):
	pattern = re.compile(r'(\d+/\d+)')
	for item in cardapio:
		if pattern.search(item):
			arrayCardapio.append(Refeicao(item, '', '', '', ''))

def adicionaLancheManha(cardapio):
	for item_array in arrayCardapio:
		data = item_array.dia
		contador = 0
		for item in cardapio:
			if item == data:
				item_array.cafe = cardapio[contador+6]
			contador = contador + 1

def adicionaAlmoco(cardapio):
	for item_array in arrayCardapio:
		data = item_array.dia
		contador = 0
		for item in cardapio:
			if item == data:
				item_array.almoco = cardapio[contador+12]
			contador = contador + 1			

def adicionaLancheTarde(cardapio):
	for item_array in arrayCardapio:
		data = item_array.dia
		contador = 0
		for item in cardapio:
			if item == data:
				item_array.lanche = cardapio[contador+18]
			contador = contador + 1						

def adicionaJanta(cardapio):
	for item_array in arrayCardapio:
		data = item_array.dia
		contador = 0
		for item in cardapio:
			if item == data:
				item_array.janta = cardapio[contador+24]
			contador = contador + 1									

def formataEmSemanas():
	verificou = False
	now = datetime.datetime.now()
	ano = now.year
	semana = 1
	contador = 1
	semana1 = []
	semana2 = []
	semana3 = []
	semana4 = []
	semana5 = []
	for item in arrayCardapio:
		dia = item.dia[0:2]
		mes = item.dia[3:5]
		data = datetime.date(ano, int(mes), int(dia))
		if semana == 1:
			brancos = data.weekday()
			if brancos == 0:
				verificou = True
			if not verificou:
				for x in range(0, brancos):
					semana1.append(Refeicao('', '', '', '', ''))
					contador = contador + 1
			semana1.append(item)
		elif semana == 2:
			semana2.append(item)
		elif semana == 3:
			semana3.append(item)
		elif semana == 4:
			semana4.append(item)
		else:
			semana5.append(item)
		contador = contador + 1
		if contador > 5:
			contador = 1
			semana = semana + 1
	final.append(semana1)
	final.append(semana2)
	final.append(semana3)
	final.append(semana4)
	final.append(semana5)

def formataTexto(texto):
	texto = texto.strip()
	if texto.find(',') < 0:
		texto += ', '
	#texto = texto.replace(',,', '')
	return texto	

def trataCelula(celula):
	retorno = ''

	for paragraph in celula.paragraphs:
		texto = formataTexto(paragraph.text)
		retorno += texto
	if retorno.endswith(', '):
		retorno = retorno[:-2]
	pattern = re.compile(r'(\w+,)(\w+)')
	substituido = pattern.sub(r'\1 \2', retorno)
	return substituido

doc = docx.Document('cardapio_outubro.docx')
fullText = []
refeicoes = []
contador = 1
multiplo = 1
conta_dias = 0
for table in doc.tables:
	for row in table.rows:
		for cell in row.cells:
			fullText.append(trataCelula(cell))

def create_index_html():
    fname = "outubro.html"
    mes = "outubro"
    aniversario = "26/10 (sexta-feira)"
    lanche = "17/10 (quarta-feira)"

    context = {
        'mes': mes,
        'aniversario': aniversario,
        'lanche': lanche,
        'cardapio': final
    }

    with open(fname, 'w') as f:
        html = render_template('index.html', context)
        f.write(html)

def main():
	adicionaDias(fullText)
	adicionaLancheManha(fullText)
	adicionaAlmoco(fullText)
	adicionaLancheTarde(fullText)
	adicionaJanta(fullText)
	formataEmSemanas()
	create_index_html()
  
if __name__ == "__main__":
    main()

#print(final[1][0].dia)
#print(final[1][0].cafe)